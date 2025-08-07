"""Advanced document processing service with hierarchical chunking and contextual retrieval"""

import os
import hashlib
import tempfile
from typing import List, Dict, Any, Tuple, Optional
import PyPDF2
from docx import Document
import logging
import httpx
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re
from dataclasses import dataclass
import asyncio
from concurrent.futures import ThreadPoolExecutor
import spacy
from collections import defaultdict
import pymupdf4llm

from config import settings
from vector_service import VectorService
from llm_service import LLMService
from database import DatabaseManager
from knowledge_graph_service import KnowledgeGraphService

logger = logging.getLogger(__name__)

# Load spaCy model for semantic chunking (optional)
try:
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except OSError:
    nlp = None
    SPACY_AVAILABLE = False
    logger.info("spaCy model not available, using basic chunking")

@dataclass
class DocumentChunk:
    content: str
    chunk_id: int
    parent_id: Optional[str]
    chunk_type: str  # 'parent' or 'child'
    metadata: Dict[str, Any]
    section_title: Optional[str] = None
    entities: List[str] = None
    semantic_density: float = 0.0
    context_window: str = ""

class AdvancedDocumentService:
    
    async def extract_with_pymupdf4llm(self, document_url: str) -> str:
        """Extract text from document URL using pymupdf4llm - main interface for advanced RAG"""
        temp_file = None
        try:
            # Download document
            temp_file = await self.download_document(document_url)
            
            # Extract text using pymupdf4llm if PDF, otherwise use appropriate extractor
            if temp_file.endswith('.pdf'):
                logger.info(f"Using pymupdf4llm for advanced PDF text extraction: {document_url}")
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_pdf, temp_file
                )
            elif temp_file.endswith('.docx'):
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_docx, temp_file
                )
            else:
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_email, temp_file
                )
            
            if not text.strip():
                raise ValueError("No text could be extracted from document")
            
            logger.info(f"Successfully extracted {len(text)} characters from {document_url}")
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from {document_url}: {e}")
            raise
        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)
    def __init__(self):
        self.vector_service = VectorService()
        self.llm_service = LLMService()
        self.kg_service = KnowledgeGraphService()
        
        # Thread pool for async operations with proper cleanup
        self.executor = ThreadPoolExecutor(max_workers=min(settings.max_workers, 2))  # Limit to 2 workers to prevent resource issues
        
        # Enhanced hierarchical text splitters
        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.parent_chunk_size,
            chunk_overlap=settings.chunk_overlap * 2,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        
        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.child_chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Enhanced section patterns for insurance documents
        self.section_patterns = [
            r'^\d+\.\s+([A-Z][^\n]+)',  # 1. SECTION TITLE
            r'^[A-Z][A-Z\s]{3,}:',      # SECTION TITLE:
            r'^\([a-z]\)\s+([A-Z][^\n]+)', # (a) Subsection
            r'^Article\s+\d+[:\.]?\s*([A-Z][^\n]*)',           # Article N
            r'^Section\s+\d+[:\.]?\s*([A-Z][^\n]*)',           # Section N
            r'^Chapter\s+\d+[:\.]?\s*([A-Z][^\n]*)',           # Chapter N
            r'^Part\s+[IVX]+[:\.]?\s*([A-Z][^\n]*)',           # Part I, II, etc.
            r'^\d+\.\d+\s+([A-Z][^\n]+)',  # 1.1 Subsection
            r'^[A-Z]\.\s+([A-Z][^\n]+)',   # A. Section
        ]
        
        # Insurance-specific entity patterns
        self.entity_patterns = {
            'amounts': r'(?:Rs\.?\s*|INR\s*|rupees?\s*)[\d,]+(?:\.\d+)?(?:\s*(?:lakh|crore))?|sum\s+insured\s+of\s+Rs\.?\s*[\d,]+',
            'percentages': r'\d+(?:\.\d+)?\s*%(?:\s*of\s+sum\s+insured)?',
            'time_periods': r'\d+\s*(?:days?|months?|years?|hrs?|hours?|weeks?)|waiting\s+period|cooling\s+period',
            'medical_procedures': r'\b(?:surgery|operation|treatment|procedure|therapy|consultation|diagnosis|screening|test|examination|biopsy|endoscopy|angioplasty|transplant|chemotherapy|radiotherapy|dialysis|physiotherapy)\b',
            'medical_conditions': r'\b(?:cancer|diabetes|hypertension|cardiac|stroke|pneumonia|appendicitis|hernia|cataract|kidney\s+stone|heart\s+attack|bypass|angioplasty)\b',
            'policy_terms': r'\b(?:premium|deductible|copay|co-payment|coverage|exclusion|benefit|claim|policy|sum\s+insured|waiting\s+period|pre-existing|network\s+hospital|cashless|reimbursement|maternity|dental|opd)\b',
            'medical_specialties': r'\b(?:cardiology|neurology|orthopedic|oncology|gynecology|pediatric|dental|ophthalmology|ENT|dermatology|urology|nephrology)\b'
        }
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file using pymupdf4llm for superior extraction"""
        try:
            logger.info(f"Extracting text from PDF using pymupdf4llm: {pdf_path}")
            
            # Use pymupdf4llm for advanced PDF text extraction
            # This provides better text extraction with layout preservation
            markdown_text = pymupdf4llm.to_markdown(pdf_path)
            
            # Convert markdown to clean text while preserving structure
            text = self._process_markdown_to_text(markdown_text)
            
            if not text.strip():
                logger.warning("pymupdf4llm returned empty text, falling back to PyPDF2")
                return self._fallback_pdf_extraction(pdf_path)
            
            logger.info(f"Successfully extracted {len(text)} characters using pymupdf4llm")
            return text
            
        except Exception as e:
            logger.error(f"Error using pymupdf4llm for {pdf_path}: {e}")
            logger.info("Falling back to PyPDF2 extraction")
            return self._fallback_pdf_extraction(pdf_path)
    
    def _process_markdown_to_text(self, markdown_text: str) -> str:
        """Convert pymupdf4llm markdown output to clean text with structure preservation"""
        try:
            # Remove markdown formatting while preserving structure
            text = markdown_text
            
            # Remove markdown headers but keep the text
            text = re.sub(r'^#{1,6}\s*(.+)$', r'\1', text, flags=re.MULTILINE)
            
            # Remove markdown emphasis but keep the text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italic
            text = re.sub(r'__(.+?)__', r'\1', text)      # Bold
            text = re.sub(r'_(.+?)_', r'\1', text)        # Italic
            
            # Clean up extra whitespace while preserving paragraph breaks
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple line breaks to double
            text = re.sub(r'[ \t]+', ' ', text)           # Multiple spaces to single
            
            # Remove markdown links but keep the text
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            
            # Remove markdown code blocks
            text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            return text.strip()
            
        except Exception as e:
            logger.warning(f"Error processing markdown text: {e}")
            return markdown_text
    
    def _fallback_pdf_extraction(self, pdf_path: str) -> str:
        """Fallback PDF extraction using PyPDF2"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                logger.info(f"Processing PDF with PyPDF2: {total_pages} pages")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += f"\n\n--- PAGE {i+1} ---\n\n" + page_text
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {i+1}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error reading PDF {pdf_path} with PyPDF2: {e}")
            raise
        return text
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from Word document with structure preservation"""
        text = ""
        try:
            doc = Document(docx_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Preserve paragraph structure
                    text += paragraph.text + "\n"
            
            # Extract tables if present
            for table in doc.tables:
                text += "\n--- TABLE ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "--- END TABLE ---\n"
                
        except Exception as e:
            logger.error(f"Error reading DOCX {docx_path}: {e}")
            raise
        return text
    
    def extract_text_from_email(self, email_path: str) -> str:
        """Extract text from email file"""
        text = ""
        try:
            with open(email_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            logger.error(f"Error reading email {email_path}: {e}")
            raise
        return text
    
    async def download_document(self, url: str) -> str:
        """Download document from URL and save to temporary file"""
        try:
            timeout = httpx.Timeout(30.0, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Determine file type from URL or content
                suffix = ".pdf"  # Default to PDF
                if url.lower().endswith('.docx'):
                    suffix = ".docx"
                elif url.lower().endswith('.txt'):
                    suffix = ".txt"
                elif url.lower().endswith('.eml'):
                    suffix = ".eml"
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                    tmp_file.write(response.content)
                    return tmp_file.name
                    
        except Exception as e:
            logger.error(f"Failed to download document from {url}: {e}")
            raise
    
    def _extract_semantic_sections(self, text: str) -> List[Tuple[str, str, float]]:
        """Extract document sections with semantic importance scoring"""
        sections = []
        current_section = "Introduction"
        current_content = ""
        current_importance = 0.5
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line matches a section pattern
            section_match = None
            importance_score = 0.5
            
            for i, pattern in enumerate(self.section_patterns):
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    section_match = match.group(0) if not match.groups() else match.group(1)
                    # Higher importance for earlier patterns (more specific)
                    importance_score = 1.0 - (i * 0.1)
                    break
            
            # Boost importance for insurance-specific sections
            if section_match and any(keyword in section_match.lower() for keyword in 
                                   ['coverage', 'exclusion', 'benefit', 'claim', 'premium', 'waiting', 'eligibility']):
                importance_score += 0.2
            
            if section_match:
                # Save previous section with importance score
                if current_content.strip():
                    sections.append((current_section, current_content.strip(), current_importance))
                
                # Start new section
                current_section = section_match.strip()
                current_content = line + '\n'
                current_importance = min(importance_score, 1.0)
            else:
                current_content += line + '\n'
        
        # Add final section
        if current_content.strip():
            sections.append((current_section, current_content.strip(), current_importance))
        
        # If no sections found, treat entire document as one section
        if len(sections) == 0:
            sections.append(("Document", text, 0.5))
        
        return sections
    
    def _calculate_semantic_density(self, text: str) -> float:
        """Calculate semantic density of text chunk"""
        if not SPACY_AVAILABLE:
            # Fallback to simple heuristics
            sentences = len(re.split(r'[.!?]+', text))
            words = len(text.split())
            entities = sum(len(re.findall(pattern, text, re.IGNORECASE)) 
                          for pattern in self.entity_patterns.values())
            
            # Simple density calculation
            if words == 0:
                return 0.0
            return min((entities * 2 + sentences) / words, 1.0)
        
        # Advanced semantic analysis with spaCy
        try:
            doc = nlp(text[:1000])  # Limit text length for processing
            
            # Count important linguistic features
            entities = len(doc.ents)
            noun_phrases = len(list(doc.noun_chunks))
            sentences = len(list(doc.sents))
            tokens = len([token for token in doc if not token.is_stop and not token.is_punct])
            
            # Calculate density score
            if tokens == 0:
                return 0.0
            
            density = (entities * 3 + noun_phrases * 2 + sentences) / tokens
            return min(density, 1.0)
            
        except Exception as e:
            logger.warning(f"Semantic density calculation failed: {e}")
            return 0.5
    
    def create_enhanced_hierarchical_chunks(self, text: str, source: str) -> Tuple[List[DocumentChunk], Dict[str, str]]:
        """Create hierarchical chunks with enhanced contextual information"""
        # Extract sections with importance scoring
        sections = self._extract_semantic_sections(text)
        
        all_chunks = []
        parent_chunks_dict = {}
        chunk_counter = 0
        
        for section_title, section_text, importance in sections:
            # Create parent chunks for each section
            parent_texts = self.parent_splitter.split_text(section_text)
            
            for parent_idx, parent_text in enumerate(parent_texts):
                parent_id = f"{source}_parent_{chunk_counter}"
                parent_chunks_dict[parent_id] = parent_text
                
                # Calculate semantic density
                semantic_density = self._calculate_semantic_density(parent_text)
                
                # Extract entities
                entities = self._extract_enhanced_entities(parent_text)
                
                # Create context window (surrounding text)
                context_start = max(0, len(section_text) // 2 - 200)
                context_end = min(len(section_text), len(section_text) // 2 + 200)
                context_window = section_text[context_start:context_end]
                
                # Create parent chunk
                parent_chunk = DocumentChunk(
                    content=parent_text,
                    chunk_id=chunk_counter,
                    parent_id=None,
                    chunk_type='parent',
                    section_title=section_title,
                    entities=entities,
                    semantic_density=semantic_density * importance,  # Boost by section importance
                    context_window=context_window,
                    metadata={
                        'source': source,
                        'section': section_title,
                        'chunk_type': 'parent',
                        'importance_score': importance,
                        'semantic_density': semantic_density,
                        'entity_count': len(entities)
                    }
                )
                all_chunks.append(parent_chunk)
                chunk_counter += 1
                
                # Create child chunks from parent with semantic chunking
                if settings.use_semantic_chunking and SPACY_AVAILABLE:
                    child_chunks = self._create_semantic_child_chunks(parent_text, parent_id, section_title, source, chunk_counter)
                else:
                    child_chunks = self._create_traditional_child_chunks(parent_text, parent_id, section_title, source, chunk_counter)
                
                all_chunks.extend(child_chunks)
                chunk_counter += len(child_chunks)
        
        return all_chunks, parent_chunks_dict
    
    def _create_semantic_child_chunks(self, parent_text: str, parent_id: str, section_title: str, source: str, start_counter: int) -> List[DocumentChunk]:
        """Create child chunks using semantic boundaries"""
        child_chunks = []
        
        try:
            doc = nlp(parent_text)
            sentences = list(doc.sents)
            
            current_chunk = ""
            current_entities = []
            chunk_counter = start_counter
            
            for sentence in sentences:
                sentence_text = sentence.text.strip()
                
                # Check if adding this sentence would exceed chunk size
                if len(current_chunk + sentence_text) > settings.child_chunk_size and current_chunk:
                    # Create chunk from current content
                    if len(current_chunk.strip()) >= 50:  # Minimum chunk size
                        entities = self._extract_enhanced_entities(current_chunk)
                        semantic_density = self._calculate_semantic_density(current_chunk)
                        
                        child_chunk = DocumentChunk(
                            content=current_chunk.strip(),
                            chunk_id=chunk_counter,
                            parent_id=parent_id,
                            chunk_type='child',
                            section_title=section_title,
                            entities=entities,
                            semantic_density=semantic_density,
                            context_window=parent_text[:200] + "..." + parent_text[-200:],
                            metadata={
                                'source': source,
                                'section': section_title,
                                'parent_id': parent_id,
                                'chunk_type': 'child',
                                'semantic_density': semantic_density,
                                'entity_count': len(entities)
                            }
                        )
                        child_chunks.append(child_chunk)
                        chunk_counter += 1
                    
                    # Start new chunk
                    current_chunk = sentence_text + " "
                else:
                    current_chunk += sentence_text + " "
            
            # Add final chunk
            if len(current_chunk.strip()) >= 50:
                entities = self._extract_enhanced_entities(current_chunk)
                semantic_density = self._calculate_semantic_density(current_chunk)
                
                child_chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=chunk_counter,
                    parent_id=parent_id,
                    chunk_type='child',
                    section_title=section_title,
                    entities=entities,
                    semantic_density=semantic_density,
                    context_window=parent_text[:200] + "..." + parent_text[-200:],
                    metadata={
                        'source': source,
                        'section': section_title,
                        'parent_id': parent_id,
                        'chunk_type': 'child',
                        'semantic_density': semantic_density,
                        'entity_count': len(entities)
                    }
                )
                child_chunks.append(child_chunk)
                
        except Exception as e:
            logger.warning(f"Semantic chunking failed, falling back to traditional: {e}")
            return self._create_traditional_child_chunks(parent_text, parent_id, section_title, source, start_counter)
        
        return child_chunks
    
    def _create_traditional_child_chunks(self, parent_text: str, parent_id: str, section_title: str, source: str, start_counter: int) -> List[DocumentChunk]:
        """Create child chunks using traditional text splitting"""
        child_chunks = []
        child_texts = self.child_splitter.split_text(parent_text)
        chunk_counter = start_counter
        
        for child_text in child_texts:
            if len(child_text.strip()) < 50:  # Skip very small chunks
                continue
            
            entities = self._extract_enhanced_entities(child_text)
            semantic_density = self._calculate_semantic_density(child_text)
            
            child_chunk = DocumentChunk(
                content=child_text,
                chunk_id=chunk_counter,
                parent_id=parent_id,
                chunk_type='child',
                section_title=section_title,
                entities=entities,
                semantic_density=semantic_density,
                context_window=parent_text[:200] + "..." + parent_text[-200:],
                metadata={
                    'source': source,
                    'section': section_title,
                    'parent_id': parent_id,
                    'chunk_type': 'child',
                    'semantic_density': semantic_density,
                    'entity_count': len(entities)
                }
            )
            child_chunks.append(child_chunk)
            chunk_counter += 1
        
        return child_chunks
    
    def _extract_enhanced_entities(self, text: str) -> List[str]:
        """Extract key entities from text chunk with enhanced patterns"""
        entities = []
        text_lower = text.lower()
        
        for entity_type, pattern in self.entity_patterns.items():
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    match = ' '.join(match)
                entities.append(match.strip())
        
        # Use spaCy for additional entity extraction if available
        if SPACY_AVAILABLE:
            try:
                doc = nlp(text[:1000])  # Limit text length
                for ent in doc.ents:
                    if ent.label_ in ['MONEY', 'PERCENT', 'DATE', 'TIME', 'ORG', 'PERSON']:
                        entities.append(ent.text.lower())
            except Exception as e:
                logger.debug(f"spaCy entity extraction failed: {e}")
        
        return list(set(entities))  # Remove duplicates
    
    async def process_document_from_url(self, url: str) -> Dict[str, Any]:
        """Process document with enhanced chunking and parallel processing"""
        temp_file = None
        try:
            # Download document
            temp_file = await self.download_document(url)
            
            # Extract text based on file type in thread pool
            if temp_file.endswith('.pdf'):
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_pdf, temp_file
                )
                file_type = "PDF (pymupdf4llm)"
            elif temp_file.endswith('.docx'):
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_docx, temp_file
                )
                file_type = "DOCX"
            else:
                text = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.extract_text_from_email, temp_file
                )
                file_type = "Text/Email"
            
            if not text.strip():
                raise ValueError("No text could be extracted from document")
            
            # Get file info
            file_size = os.path.getsize(temp_file)
            filename = os.path.basename(url).split('?')[0]
            content_hash = hashlib.md5(text.encode()).hexdigest()
            
            # Check if document already processed
            with DatabaseManager() as db:
                existing_doc = db.get_document_by_hash(content_hash)
                if existing_doc:
                    logger.info(f"Document already processed: {filename}")
                    return {
                        'document_id': existing_doc.id,
                        'chunks_processed': existing_doc.total_chunks,
                        'cached': True
                    }
            
            # Create enhanced hierarchical chunks
            chunks, parent_chunks_dict = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.create_enhanced_hierarchical_chunks, text, filename
            )
            
            # Prepare documents for vector storage with enhanced metadata
            documents = []
            for chunk in chunks:
                doc_data = {
                    'content': chunk.content,
                    'source': filename,
                    'chunk_id': chunk.chunk_id,
                    'chunk_type': chunk.chunk_type,
                    'metadata': {
                        **chunk.metadata,
                        'section_title': chunk.section_title or '',
                        'entities': chunk.entities or [],
                        'semantic_density': chunk.semantic_density,
                        'context_window': chunk.context_window
                    }
                }
                
                # Only add parent_id if it exists and is not None
                if chunk.parent_id is not None:
                    doc_data['parent_id'] = chunk.parent_id
                
                documents.append(doc_data)
            
            # Store in vector database with parent chunks (async)
            vector_ids = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.vector_service.upsert_documents, documents, parent_chunks_dict
            )
            
            # Process for knowledge graph if enabled (async)
            if self.kg_service.enabled:
                try:
                    await asyncio.get_event_loop().run_in_executor(
                        self.executor, self.kg_service.process_document_for_kg, text, filename
                    )
                    logger.info(f"Knowledge graph processing completed for {filename}")
                except Exception as e:
                    logger.warning(f"Knowledge graph processing failed: {e}")
            
            # Store in database
            with DatabaseManager() as db:
                doc_record = db.save_document(
                    url=url,
                    filename=filename,
                    content_hash=content_hash,
                    total_chunks=len(documents),
                    file_size=file_size,
                    file_type=file_type
                )
                
                # Save document chunks with enhanced metadata
                for doc, vector_id in zip(documents, vector_ids):
                    db.save_document_chunk(
                        document_id=doc_record.id,
                        chunk_index=doc['chunk_id'],
                        content=doc['content'],
                        pinecone_id=vector_id,
                        chunk_metadata=doc['metadata']
                    )
                
                document_id = doc_record.id
            
            return {
                'document_id': document_id,
                'chunks_processed': len(documents),
                'cached': False,
                'processing_stats': {
                    'total_chunks': len(documents),
                    'parent_chunks': len([c for c in chunks if c.chunk_type == 'parent']),
                    'child_chunks': len([c for c in chunks if c.chunk_type == 'child']),
                    'avg_semantic_density': sum(c.semantic_density for c in chunks) / len(chunks) if chunks else 0,
                    'total_entities': sum(len(c.entities or []) for c in chunks)
                }
            }
            
        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)
    
    def search_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for relevant document chunks with enhanced ranking"""
        results = self.vector_service.search_similar(query=query, top_k=top_k)
        
        # Boost results based on semantic density and entity count
        for result in results:
            metadata = result.get('metadata', {})
            semantic_density = metadata.get('semantic_density', 0.5)
            entity_count = metadata.get('entity_count', 0)
            
            # Calculate boost factor
            boost = 1.0 + (semantic_density * 0.2) + (min(entity_count, 10) * 0.05)
            result['relevance_score'] = result.get('relevance_score', 0) * boost
        
        # Re-sort by boosted scores
        results.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
        
        return results
    
    async def process_question(self, question: str, document_url: str) -> str:
        """Process question with advanced query decomposition and multi-step reasoning"""
        try:
            # Process document if not already done
            doc_info = await self.process_document_from_url(document_url)
            
            # Decompose query for advanced processing
            query_plan = self.llm_service.decompose_query(question)
            
            # Parse the question with enhanced context
            parsed_query = self.llm_service.parse_query(question)
            
            # Generate targeted search queries
            search_queries = self.llm_service.generate_search_queries(parsed_query)
            
            # Perform advanced retrieval
            relevant_documents = await self._advanced_retrieval(search_queries, parsed_query)
            
            # Generate answer using multi-step reasoning if enabled
            if settings.use_query_decomposition:
                return self.llm_service.generate_multi_step_answer(query_plan, relevant_documents)
            else:
                return self.llm_service.generate_direct_answer(question, relevant_documents)
            
        except Exception as e:
            logger.error(f"Error processing question '{question}': {e}")
            return f"Unable to process question due to error: {str(e)}"
    
    async def _advanced_retrieval(self, search_queries: List[str], 
                                parsed_query: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Advanced retrieval with query-specific optimization and parallel processing"""
        all_results = []
        
        # Retrieve with different strategies based on query type
        query_type = parsed_query.get('query_type', 'general')
        entities = parsed_query.get('entities', [])
        
        # Parallel base retrieval for all search queries
        retrieval_tasks = []
        for search_query in search_queries:
            task = asyncio.get_event_loop().run_in_executor(
                self.executor, self.search_documents, search_query, settings.final_top_k
            )
            retrieval_tasks.append(task)
        
        # Execute all retrieval tasks in parallel
        if retrieval_tasks:
            parallel_results = await asyncio.gather(*retrieval_tasks, return_exceptions=True)
            for result in parallel_results:
                if isinstance(result, list):
                    all_results.extend(result)
                elif isinstance(result, Exception):
                    logger.warning(f"Retrieval task failed: {result}")
        
        # Entity-specific retrieval with knowledge graph enhancement
        kg_enhanced_queries = []
        if self.kg_service.enabled and entities:
            try:
                kg_enhanced_queries = await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.kg_service.enhance_query_with_kg, entities
                )
                logger.info(f"Generated {len(kg_enhanced_queries)} KG-enhanced queries")
            except Exception as e:
                logger.warning(f"KG query enhancement failed: {e}")
        
        # Use KG-enhanced queries if available
        if kg_enhanced_queries:
            kg_tasks = []
            for kg_query in kg_enhanced_queries[:5]:  # Limit KG queries
                task = asyncio.get_event_loop().run_in_executor(
                    self.executor, self.search_documents, kg_query, 3
                )
                kg_tasks.append(task)
            
            if kg_tasks:
                kg_results = await asyncio.gather(*kg_tasks, return_exceptions=True)
                for result in kg_results:
                    if isinstance(result, list):
                        all_results.extend(result)
        
        # Query-type specific retrieval
        if query_type in ['waiting_period', 'eligibility']:
            waiting_results = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.search_documents, "waiting period eligibility requirements", 3
            )
            all_results.extend(waiting_results)
        elif query_type == 'exclusion':
            exclusion_results = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.search_documents, "exclusions limitations not covered", 3
            )
            all_results.extend(exclusion_results)
        
        # Deduplicate and rank
        return self._deduplicate_and_rank_results(all_results)
    
    def _deduplicate_and_rank_results(self, results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicates and rank results by multiple criteria"""
        unique_results = {}
        
        for result in results:
            content_hash = hashlib.md5(result['content'].encode()).hexdigest()
            
            if content_hash not in unique_results:
                unique_results[content_hash] = result
            else:
                # Boost score if found multiple times
                existing = unique_results[content_hash]
                boost_score = existing.get('final_score', existing.get('relevance_score', 0)) * 1.1
                existing['final_score'] = boost_score
        
        # Sort by final score or relevance score
        sorted_results = sorted(
            unique_results.values(),
            key=lambda x: x.get('final_score', x.get('relevance_score', 0)),
            reverse=True
        )
        
        return sorted_results[:settings.final_top_k]
    
    async def process_multiple_questions(self, questions: List[str], document_url: str) -> List[str]:
        """Process multiple questions with optimized document processing and parallel execution"""
        answers = []
        
        # Process document once
        try:
            doc_info = await self.process_document_from_url(document_url)
            logger.info(f"Document processed: {doc_info['chunks_processed']} chunks, cached: {doc_info['cached']}")
            if 'processing_stats' in doc_info:
                logger.info(f"Processing stats: {doc_info['processing_stats']}")
        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            return [f"Error processing document: {str(e)}" for _ in questions]
        
        # Process questions in parallel batches to avoid overwhelming the system
        batch_size = min(3, len(questions))  # Process 3 questions at a time
        
        for i in range(0, len(questions), batch_size):
            batch = questions[i:i + batch_size]
            batch_tasks = []
            
            for j, question in enumerate(batch):
                logger.info(f"Processing question {i+j+1}/{len(questions)}: {question[:100]}...")
                answer = await self.process_question(question, document_url)
                batch_tasks.append(answer)
            
            # Execute batch in parallel
            try:
                batch_answers = await asyncio.gather(*batch_tasks, return_exceptions=True)
                
                for answer in batch_answers:
                    if isinstance(answer, Exception):
                        logger.error(f"Error processing question: {answer}")
                        answers.append(f"Unable to process question due to error: {str(answer)}")
                    else:
                        answers.append(answer)
                        
            except Exception as e:
                logger.error(f"Batch processing failed: {e}")
                answers.extend([f"Unable to process question due to error: {str(e)}" for _ in batch])
        
        return answers
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Get system statistics"""
        try:
            vector_stats = self.vector_service.get_index_stats()
            
            with DatabaseManager() as db:
                performance_stats = db.get_performance_stats(24)
            
            return {
                'vector_stats': vector_stats,
                'performance_stats': performance_stats,
                'llm_healthy': self.llm_service.health_check(),
                'vector_healthy': self.vector_service.health_check(),
                'spacy_available': SPACY_AVAILABLE,
                'semantic_chunking_enabled': settings.use_semantic_chunking,
                'async_processing_enabled': settings.use_async_processing
            }
        except Exception:
            return {}
    
    def health_check(self) -> Dict[str, bool]:
        """Check health of all services"""
        health_status = {
            'llm': self.llm_service.health_check(),
            'vector': self.vector_service.health_check(),
            'database': True,
            'spacy': SPACY_AVAILABLE
        }
        
        # Add knowledge graph health if enabled
        if self.kg_service.enabled:
            health_status['knowledge_graph'] = self.kg_service.health_check()
        
        return health_status
    
    def cleanup(self):
        """Clean up resources"""
        try:
            if hasattr(self, 'executor') and self.executor:
                self.executor.shutdown(wait=True)
                logger.info("Thread pool executor shut down successfully")
        except Exception as e:
            logger.warning(f"Error during cleanup: {e}")

# For backward compatibility, alias the new class
DocumentService = AdvancedDocumentService